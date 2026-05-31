#!/usr/bin/env python3
"""Fix all inaccuracies in RoadSOS_Submission.docx.

Performs text-level replacements across all paragraphs and table cells
in the DOCX to align it with the current codebase.
"""
import re
from docx import Document
from pathlib import Path

DOCX_PATH = Path(__file__).resolve().parent.parent / "docs" / "RoadSOS_Submission.docx"

# ── Factual corrections ─────────────────────────────────────────────────
REPLACEMENTS = [
    # 1. Bundled facility count
    ("818 entries", "938 entries"),

    # 2. Overpass strategy (multiple phrasings found in the doc)
    ("3-mirror Overpass with exponential back-off",
     "3 Overpass mirrors raced concurrently (first healthy response wins)"),
    ("Three mirrors with exponential-backoff retry",
     "Three mirrors raced concurrently (first healthy response wins)"),
    ("Three-mirror Overpass retry with exponential backoff",
     "Three Overpass mirrors raced concurrently (first healthy response wins)"),
    ("3 mirrors, 12\u2009s timeout each",   # thin space variant
     "3 mirrors raced concurrently, 10\u2009s per-mirror timeout"),
    ("3 mirrors, 12 s timeout each",
     "3 mirrors raced concurrently, 10 s per-mirror timeout"),

    # 3. Auto-expand radius
    ("5 km → 10 km → 20 km",  "8 km → 25 km"),
    ("5\u2009km → 10\u2009km → 20\u2009km", "8\u2009km → 25\u2009km"),  # thin space
    ("5 km $\\to$ 10 km $\\to$ 20 km", "8 km → 25 km"),

    # 4. localStorage TTL
    ("24 h localStorage cache", "7-day localStorage cache"),
    ("24\u2009h localStorage cache", "7-day localStorage cache"),

    # 5. Gemini model version (only fix the wrong "2.0")
    ("Gemini 2.0 Flash", "Gemini 2.5 Flash"),
    ("Google Gemini 2.0 Flash", "Google Gemini 2.5 Flash"),

    # 6. Overpass per-attempt budget
    ("12 s timeout each", "10 s per-mirror timeout"),
    ("12\u2009s timeout each", "10\u2009s per-mirror timeout"),

    # 7. hardcoded mock → national emergency numbers banner
    ("hardcoded mock", "national emergency-numbers banner"),

    # ── Unicode escape sequences from build_submission_tex.py ──
    ("[U+2014]", "\u2014"),   # em-dash  —
    ("[U+2013]", "\u2013"),   # en-dash  –
    ("[U+2192]", "\u2192"),   # arrow    →
    ("[U+00B7]", "\u00B7"),   # middle dot ·
    ("[U+2248]", "\u2248"),   # approx   ≈
    ("[U+00D7]", "\u00D7"),   # times    ×
    ("[U+00A0]", "\u00A0"),   # nbsp
    ("[U+2019]", "\u2019"),   # right single quote '
    ("[U+201C]", "\u201C"),   # left double quote "
    ("[U+201D]", "\u201D"),   # right double quote "
    ("[U+2026]", "\u2026"),   # ellipsis …
    ("[U+2264]", "\u2264"),   # ≤
    ("[U+2265]", "\u2265"),   # ≥
]


def apply_replacements(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def fix_paragraph(para):
    """Apply replacements to every run in a paragraph."""
    full = para.text
    if not full:
        return False

    new_full = apply_replacements(full)
    if new_full == full:
        return False

    # If the paragraph has a single run, simple replacement
    if len(para.runs) == 1:
        para.runs[0].text = new_full
        return True

    # Multi-run: rebuild — put all text in the first run, clear the rest
    # (preserves the first run's formatting)
    para.runs[0].text = new_full
    for run in para.runs[1:]:
        run.text = ""
    return True


def main():
    doc = Document(str(DOCX_PATH))
    changes = 0

    # Fix paragraphs in the body
    for para in doc.paragraphs:
        if fix_paragraph(para):
            changes += 1

    # Fix paragraphs inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if fix_paragraph(para):
                        changes += 1

    doc.save(str(DOCX_PATH))
    print(f"Applied replacements across {changes} paragraphs/cells in {DOCX_PATH.name}")


if __name__ == "__main__":
    main()
